"""
docx_style.py — Apply institutional memo styling to a python-docx Document.

Produces IC memos / LP letters with the conservative institutional look:
- Calibri 11pt body (or Garamond 11pt traditional)
- Numbered hierarchy (1., 1.1, 1.1.1)
- Sans-serif headings or serif-on-serif option
- Minimal table borders (top/bottom only, plus rule under header row)
- 1" margins, page numbers, optional confidentiality marker
- Right-aligned numerics in tables

Usage:
    from docx import Document
    from docx_style import apply_memo_styles, add_heading, add_para, add_table, add_cover_page

    doc = Document()
    apply_memo_styles(doc, theme="modern")  # or "traditional"
    add_cover_page(doc, title="Disposition Recommendation — Marina Apartments",
                   recipient="Investment Committee", preparer="Asset Management",
                   date_str="May 11, 2026", confidentiality=True)
    add_heading(doc, "1. Recommendation", level=1)
    add_para(doc, "Sell at $84.5M (5.1% cap), generating net proceeds of $39.8M ...")
    add_table(doc, headers=["Metric", "UW", "Actual"], rows=[["NOI", "$4.2M", "$4.5M"]])
    doc.save("memo.docx")

CLI smoke-test:
    python docx_style.py --demo output.docx
"""

from __future__ import annotations

import argparse

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor


# ---------------------------------------------------------------------------
# Palette
# ---------------------------------------------------------------------------

INK = RGBColor(0x22, 0x22, 0x22)
NAVY = RGBColor(0x1F, 0x38, 0x64)
CHARCOAL = RGBColor(0x33, 0x33, 0x33)
MID_GREY = RGBColor(0x7F, 0x7F, 0x7F)
SOFT_GREY_HEX = "D9D9D9"
PAPER_GREY_HEX = "F2F2F2"


# ---------------------------------------------------------------------------
# Theme
# ---------------------------------------------------------------------------

THEMES = {
    "modern": {
        "body_font":    "Calibri",
        "body_size":    Pt(11),
        "h1_font":      "Calibri",
        "h1_size":      Pt(14),
        "h2_font":      "Calibri",
        "h2_size":      Pt(12),
        "h3_font":      "Calibri",
        "h3_size":      Pt(11),
    },
    "traditional": {
        "body_font":    "Garamond",
        "body_size":    Pt(11),
        "h1_font":      "Garamond",
        "h1_size":      Pt(14),
        "h2_font":      "Garamond",
        "h2_size":      Pt(12),
        "h3_font":      "Garamond",
        "h3_size":      Pt(11),
    },
    "times": {
        "body_font":    "Times New Roman",
        "body_size":    Pt(11),
        "h1_font":      "Times New Roman",
        "h1_size":      Pt(14),
        "h2_font":      "Times New Roman",
        "h2_size":      Pt(12),
        "h3_font":      "Times New Roman",
        "h3_size":      Pt(11),
    },
}


# ---------------------------------------------------------------------------
# Setup
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, hex_color: str) -> None:
    """Apply a fill color to a table cell (low-level XML)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_cell_border(cell, **kwargs) -> None:
    """
    Set borders on a single cell. kwargs keys: top, bottom, left, right; values are dicts
    like {"sz": 6, "val": "single", "color": "222222"}. Default value applied to keys present.
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        spec = kwargs.get(edge)
        elt = OxmlElement(f"w:{edge}")
        if spec is None:
            elt.set(qn("w:val"), "nil")
        else:
            for k, v in spec.items():
                elt.set(qn(f"w:{k}"), str(v))
        tc_borders.append(elt)
    tc_pr.append(tc_borders)


def apply_memo_styles(doc: Document, theme: str = "modern") -> None:
    """Configure base styles, margins, and headers for a memo document."""
    if theme not in THEMES:
        raise ValueError(f"unknown theme: {theme}; choose from {list(THEMES)}")
    t = THEMES[theme]

    # Body / Normal style
    style = doc.styles["Normal"]
    font = style.font
    font.name = t["body_font"]
    font.size = t["body_size"]
    font.color.rgb = INK
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.15

    # Heading styles
    for level, key in [(1, "h1"), (2, "h2"), (3, "h3")]:
        h = doc.styles[f"Heading {level}"]
        h.font.name = t[f"{key}_font"]
        h.font.size = t[f"{key}_size"]
        h.font.bold = True
        h.font.color.rgb = NAVY if level == 1 else INK
        h.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        h.paragraph_format.space_after = Pt(4)

    # Margins on all sections
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Add page numbers to footer (right-aligned)
    _add_page_numbers(doc)


def _add_page_numbers(doc: Document) -> None:
    """Insert a page number field, right-aligned, in the primary footer."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    # Field code: PAGE
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    # Style
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MID_GREY


# ---------------------------------------------------------------------------
# Content helpers
# ---------------------------------------------------------------------------

def add_cover_page(
    doc: Document,
    title: str,
    recipient: str | None = None,
    preparer: str | None = None,
    date_str: str | None = None,
    confidentiality: bool = True,
) -> None:
    """Add a clean cover page with title, recipient, preparer, date."""
    if confidentiality:
        p = doc.add_paragraph("CONFIDENTIAL — FOR DISCUSSION PURPOSES ONLY")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = MID_GREY
            run.font.bold = True

    # Title
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = NAVY

    for _ in range(3):
        doc.add_paragraph()

    if recipient:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Prepared for: {recipient}")
        run.font.size = Pt(12)
    if preparer:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Prepared by: {preparer}")
        run.font.size = Pt(12)
    if date_str:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(date_str)
        run.font.size = Pt(12)

    # Page break to start the body on page 2
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    if level not in (1, 2, 3):
        raise ValueError("level must be 1, 2, or 3")
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.font.bold = True
    if italic:
        run.font.italic = True


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    numeric_cols: list[int] | None = None,
    total_row: bool = False,
) -> None:
    """
    Render an institutional-style table.
    - Top + bottom border on the table, single rule under the header.
    - No vertical borders, no grid fill.
    - Header row: bold, light grey fill.
    - numeric_cols: zero-indexed columns to right-align.
    - total_row: if True, the last row gets a thin top border + bold.
    """
    numeric_cols = numeric_cols or []
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table.autofit = True

    # Header
    hdr = table.rows[0]
    for i, txt in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = txt
        _set_cell_shading(cell, PAPER_GREY_HEX)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
        _set_cell_border(
            cell,
            top={"val": "single", "sz": 6, "color": "222222"},
            bottom={"val": "single", "sz": 6, "color": "222222"},
        )

    # Body rows
    for r_idx, row_data in enumerate(rows):
        is_last = r_idx == len(rows) - 1
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if c_idx in numeric_cols else WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.size = Pt(10)
                    if total_row and is_last:
                        run.font.bold = True
            # Border: bottom rule on last row; thin top on last row if total
            if total_row and is_last:
                _set_cell_border(
                    cell,
                    top={"val": "single", "sz": 6, "color": "222222"},
                    bottom={"val": "single", "sz": 6, "color": "222222"},
                )
            elif is_last:
                _set_cell_border(
                    cell,
                    bottom={"val": "single", "sz": 6, "color": "222222"},
                )


def add_source(doc: Document, text: str) -> None:
    """Italic grey source line, 9pt. Always prefixed with 'Source: '."""
    if not text.startswith("Source:") and not text.startswith("Note:"):
        text = f"Source: {text}"
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = MID_GREY


# ---------------------------------------------------------------------------
# Demo
# ---------------------------------------------------------------------------

def _build_demo(path: str) -> None:
    doc = Document()
    apply_memo_styles(doc, theme="modern")

    add_cover_page(
        doc,
        title="Disposition Recommendation\nMarina Apartments",
        recipient="Investment Committee",
        preparer="Asset Management Team",
        date_str="May 11, 2026",
        confidentiality=True,
    )

    add_heading(doc, "1. Recommendation", level=1)
    add_para(doc, "Sell Marina Apartments at $84.5M (5.10% cap rate on T-12 NOI), generating "
                  "net proceeds to the fund of $39.8M and a realized deal-level net IRR of 22.4% / "
                  "1.78x MOIC. The recommended marketing launch date is June 15, 2026.")

    add_heading(doc, "2. Rationale", level=1)
    add_heading(doc, "2.1 Stub IRR vs. Hold", level=2)
    add_para(doc, "Holding through the original UW exit (Q4 2027) generates an incremental stub IRR "
                  "of 7.8% from today's equity value, materially below the 15% hold hurdle.")

    add_table(
        doc,
        headers=["Metric", "Sell Now", "Hold to UW Exit", "Δ"],
        rows=[
            ["Today's equity value",  "$39.8M",  "$39.8M",  "—"],
            ["Interim distributions",  "—",       "$3.6M",   "+$3.6M"],
            ["Exit proceeds",          "—",       "$45.2M",  "+$45.2M"],
            ["Stub IRR",               "n/a",     "7.8%",    "—"],
            ["Hurdle",                 "—",       "15.0%",   "—"],
        ],
        numeric_cols=[1, 2, 3],
    )
    add_source(doc, "Internal pro forma; CBRE BOV dated April 28, 2026.")

    add_heading(doc, "2.2 Market Conditions", level=2)
    add_para(doc, "Submarket cap rates have compressed 35 bps over the past two quarters as 10-year "
                  "Treasury yields stabilized. Recent comparable transactions price at 4.85% – 5.30%, "
                  "with bidder depth concentrated among core-plus and 1031 buyers.")

    add_heading(doc, "3. Process", level=1)
    add_bullets(doc, [
        "Solicit BOVs from CBRE, JLL, and Newmark by May 20, 2026.",
        "Execute listing agreement by May 30, 2026.",
        "OM distribution: June 15, 2026.",
        "Call for offers: July 22, 2026.",
        "PSA execution: target August 30, 2026; close target October 31, 2026.",
    ])

    doc.save(path)


def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--demo", metavar="OUTPUT", help="Write a styled demo memo to OUTPUT.docx")
    args = ap.parse_args()
    if args.demo:
        _build_demo(args.demo)
        print(f"Wrote demo memo to {args.demo}")
    else:
        ap.print_help()


if __name__ == "__main__":
    main()
